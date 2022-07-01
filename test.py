import docx.document
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()
from docx.oxml.table import CT_Tbl, CT_Row
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
            # table = Table(child, parent)
            # for row in table.rows:
            #     for cell in row.cells:
            #         yield from iter_block_items(cell)


def insert_row(table, ix):
    tbl = table._tbl
    successor = tbl.tr_lst[ix]
    tr = tbl._new_tr()
    w = 0
    for gridCol in tbl.tblGrid.gridCol_lst:
        w += gridCol.w
        #tc = tr.add_tc()
        #tc.width = gridCol.w
    tc = tr.add_tc()
    tc.width = w
    successor.addprevious(tr)
    return table.rows[ix]


doc = Document("./TestWord/simpletableau.docx")
print(type(doc))

for i, block in enumerate(iter_block_items(doc)):
    print('i= ' + str(i) + ' ' + str(type(block)))
    if isinstance(block, Paragraph):
        print('paragraphe')
        print(block.text)
    if isinstance(block, Table):
        print('Table')
        print(len(block.rows))

        for x, row in enumerate(block.rows):
            print('x: ' + str(x))
            print(row._tr.xml)

        insert_row(block,2)

doc.save('testtable.docx')
