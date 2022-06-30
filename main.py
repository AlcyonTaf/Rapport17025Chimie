# -*- coding: utf-8 -*-
import pandas
import pandas as pd
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()

# Soit on iter sur tous le documents dans l'ordre, soit on utilise les index.
# Index tableau :
# 0 = tableau informations demandes = Pas de modif
# 1 = Tableau informations item = Pas de modif
# 2 = Tableau informations essais = Mettre status des essais et ou num page de l'essai
# 3 = Tableau de révision = on ne fait rien
# 4 = 1er essai : modifier format des tableau imbriqué =Voir details ci-dessous
# 4 + n-1 = Tableau des autres essais
# Dernier tableau : Signature = on ne fait rien

# Détails tableau imbriqué dans essais :
# ligne 4 : Tableau maintenance equipement : one ne fait rien
# ligne 7 : Tableau des caractéristiques demandés : Mise ne forme a faire
# Ligne 9 : Tableau conditions de l'essais : Mise en forme a faire
# Ligne 11 : Tableau des résultats : Mise en forme a faire
# On va faire la même mise en forme pour chacun de ces tableaux

'''
#############
# Variables #
#############
'''
nbr_tableau_fixe = 5  # PAs util, Nbr de tableau qui ne bouge pas, c'est-à-dire qui ne sont pas des tableaux d'essais
list_row_nested_tab_essais = [4, 7, 9, 11]  # Emplacement des tableaux imbriqué dans les tableaux essais

''''
 Fonction
'''

def iter_tables_with_table(table):
    for row in table.rows:
        for cell in row.cells:
            for nested_table in cell.tables:
                yield nested_table
                yield from iter_tables_with_table(nested_table)


def format_tab_essais(table):
    """
        Reçoit en entré le tableau d'essai et effectue toutes les mises en forme
    """
    print(len(table.rows))
    # les tableau imbriqué sont toujours au même position, mais parfois il ne sont pas présent
    for x in list_row_nested_tab_essais:
        # Check if table exist
        if table.cell(x, 0).tables:
            nested_table = table.cell(x, 0).tables[0]
            print("table ok")
            # Code trouvé sur stackoverflow : https://stackoverflow.com/questions/58254609/python-docx-parse-a-table-to-panda-dataframe
            df = [['' for i in range(len(nested_table.columns))] for j in range(len(nested_table.rows))]
            for y, row in enumerate(nested_table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[y][j] = cell.text

            data = pd.DataFrame(df)
            print(data)

            # Ici on a récupérer les datas de la table, on va supprimer le tableau
            nested_table._element.getparent().remove(nested_table._element)

            # Ensuite ont créé les nouveaux tableaux
            table.cell(x,0).split()




        else:
            print('pas de table')

        # print(len(table.cell(i, 0).tables))
        # print(table.cell(x,0).text)
        # print(table.cell(9,0)._tc.xml)
        # print(table.cell(7,0))


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    doc = Document("./TestWord/SansMacro.docx")
    all_paras = doc.paragraphs
    # print(len(all_paras))
    tb = doc.tables
    print('Len(tb) = ' + str(len(tb)))

    # for i, table in enumerate(tb):
    #     print('i=' + str(i))
    #     print(str(len(table.rows)))

    # On récupère les positions des tables d'essais:
    # print("nbr de tableau d'essai : " + str(len(tb) - nbr_tableau_fixe))
    list_index_tab = range(4, len(tb) - 1)
    # on boucle sur chaque tableau d'essai et on appele le functon de mise en forme
    for i in list_index_tab:
        print(" Position du tableau d'essai : " + str(i))
        format_tab_essais(tb[i])

    # data =[]
    # for i, row in enumerate(tb.rows):
    #     print(i)
    #     text = (cell.text for cell in row.cells)
    #
    #     # Establish the mapping based on the first row
    #     # headers; these will become the keys of our dictionary
    #     if i == 0:
    #         keys = tuple(text)
    #         print(keys)
    #         continue
    #
    #     # Construct a dictionary for this row, mapping
    #     # keys to values for this row
    #     row_data = dict(zip(keys, text))
    #     print(row_data)
    #     data.append(row_data)
    #
    # print(data)

    doc.save('Test.docx')
